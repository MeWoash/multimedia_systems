import numpy as np
import os
import sys
import matplotlib.pyplot as plt
import pandas as pd

import cv2

#doc generation
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

import scipy.fftpack

INPUT_DIR = os.path.normpath(os.path.join(__file__,'..','input'))
OUTPUT_DIR = os.path.normpath(os.path.join(__file__,'..','output'))


##############################################################################
####     Kompresja i dekompresja    ##########################################
##############################################################################

def encode_none(A):
    return np.copy(A)

def decode_none(A):
    return np.copy(A)

def index_same_numbers(start, data) -> list:
    maxsize = data.size
    element_to_find = data[start]
    repetitions = 1
    current_index = start + 1

    while current_index < maxsize:
        if data[current_index] == element_to_find:
            repetitions+=1
        else:
            break
        current_index += 1

    return [int(repetitions), int(current_index)]

def index_different_numbers(start:int, data:int) -> list:
    maxsize = data.size
    last_element = data[start]
    streak = 1
    current_index = start + 1

    while current_index < maxsize:
        if data[current_index] != last_element:
            streak+=1
            last_element = data[current_index]
        else:
            break
        current_index += 1

    if current_index < maxsize:
        return [int(streak-1), int(current_index-1)]
    else :
        return [int(streak), int(current_index)]

def encode_ByteRun(data: np.ndarray):
    data_copy = data.copy()
    max_size = data.size
    x=np.array([len(data.shape)])
    x=np.concatenate([x, data.shape])

    data_copy = data_copy.flatten()
    bufor = np.zeros(int(np.prod(data.shape)*2,)).astype(int)

    start = 0
    bufor_current_index = 0

    while start < max_size:
        current_element = data_copy[start]

        repetitions_tmp, start_tmp = index_same_numbers(start, data_copy)
        if repetitions_tmp > 1:
            bufor[[bufor_current_index, bufor_current_index+1]] = [-repetitions_tmp+1, current_element]
            repetitions, start = repetitions_tmp, start_tmp
            bufor_current_index += 2
        else:
            repetitions_tmp, start_tmp = index_different_numbers(start, data_copy)
            bufor[bufor_current_index : bufor_current_index+repetitions_tmp+1] = np.concatenate([np.array([repetitions_tmp-1]), data_copy[start:start_tmp]])

            repetitions, start = repetitions_tmp, start_tmp
            bufor_current_index += repetitions + 1

    bufor = np.concatenate([x, bufor[:bufor_current_index]])
    return bufor

def decode_ByteRun(data: np.ndarray):
    max_size = data.size
    shape = data[1:int(data[0]+1)].astype(int)
    bufor_size = np.prod(shape)
    bufor = np.zeros(int(bufor_size)).astype(int)

    start = int(data[0])+1
    bufor_current_index = 0

    while start<max_size:
        repetitions = data[start]

        if repetitions <= -1:
            repetitions = -repetitions + 1
            element = data[start + 1]
            seq = np.repeat(element, repetitions)

            bufor[bufor_current_index:bufor_current_index+int(repetitions)] = seq
            bufor_current_index += int(repetitions)
            start+=2
        else:
            repetitions = repetitions + 1
            bufor[bufor_current_index:bufor_current_index+repetitions] = data[start+1:start+repetitions+1]
            bufor_current_index += int(repetitions)
            start+=repetitions+1
        

    bufor=bufor.reshape(shape)

    return bufor

def encode_RLE(data: np.ndarray):
    data_copy = data.copy()
    max_size = data.size
    x=np.array([len(data.shape)])
    x=np.concatenate([x, data.shape])

    data_copy = data_copy.flatten()
    bufor = np.zeros(int(np.prod(data.shape)*2,)).astype(int)

    start = 0
    bufor_current_index = 0

    while start < max_size:
        current_element = data_copy[start]

        repetitions, start = index_same_numbers(start, data_copy)
        bufor[[bufor_current_index, bufor_current_index+1]] = [repetitions, current_element]
        bufor_current_index += 2
        

    bufor = np.concatenate([x, bufor[:bufor_current_index]])
    return bufor

def decode_RLE(data: np.ndarray):
    max_size = data.size
    shape = data[1:int(data[0]+1)].astype(int)
    bufor_size = np.sum(data[int(data[0])+1::2])
    bufor = np.zeros(int(bufor_size)).astype(int)

    start = int(data[0])+1
    bufor_current_index = 0

    while start<max_size:
        repetitions = data[start]
        element = data[start + 1]
        seq = np.repeat(element, repetitions)

        bufor[bufor_current_index:bufor_current_index+int(repetitions)] = seq
        bufor_current_index += int(repetitions)
        start+=2

    bufor=bufor.reshape(shape)

    return bufor

class data:
    def init(self):
        self.Y=None
        self.Cb=None
        self.Cr=None

def Chroma_subsampling(L,subsampling):
    B = np.copy(L)
    if subsampling == "4:2:2":
        B = B[:,::2]
    elif subsampling == "4:2:0":
        B = B[::2, ::2]
    elif subsampling == "4:1:1":
        B = B[:, ::4]
    elif subsampling == "4:1:0":
        B = B[::2, ::4]
    elif subsampling == "4:4:0":
        B = B[::2, :]
    else: #4:4:4
        pass
    return B

def Chroma_resampling(L,subsampling):
    B = np.copy(L)

    if subsampling == "4:2:2":
        B = np.repeat(B, 2, axis=1)

    elif subsampling == "4:2:0":
        B = np.repeat(B, 2, axis=1)
        B = np.repeat(B, 2, axis=0)

    elif subsampling == "4:1:1":
        B = np.repeat(B, 4, axis=1)

    elif subsampling == "4:1:0":
        B = np.repeat(B, 4, axis=1)
        B = np.repeat(B, 2, axis=0)

    elif subsampling == "4:4:0":
         B = np.repeat(B, 2, axis=0)

    else: #4:4:4
        pass

    return B

        
def frame_image_to_class(frame,subsampling):
    frame_copy = np.copy(frame)
    Frame_class = data()
    Frame_class.Y=frame_copy[:,:,0].astype(int)
    Frame_class.Cb=Chroma_subsampling(frame_copy[:,:,2].astype(int),subsampling)
    Frame_class.Cr=Chroma_subsampling(frame_copy[:,:,1].astype(int),subsampling)
    return Frame_class


def frame_layers_to_image(Y, Cr, Cb, subsampling):  
    Cb=Chroma_resampling(Cb,subsampling)
    Cr=Chroma_resampling(Cr,subsampling)
    return np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)

def compress_KeyFrame(Frame_class: data, encode):
    KeyFrame = data()
    ## TO DO
    KeyFrame.Y=encode(Frame_class.Y)
    KeyFrame.Cb=encode(Frame_class.Cb)
    KeyFrame.Cr=encode(Frame_class.Cr)
    return KeyFrame

def decompress_KeyFrame(KeyFrame: data, subsampling, decode):
    Y=decode(KeyFrame.Y)
    Cb=decode(KeyFrame.Cb)
    Cr=decode(KeyFrame.Cr)

    Decompresed_KeyFrame = data()
    Decompresed_KeyFrame.Y = Y
    Decompresed_KeyFrame.Cr = Cb
    Decompresed_KeyFrame.Cb = Cr
    ## TO DO 
    frame_image=frame_layers_to_image(Y,Cr,Cb,subsampling)
    return frame_image, Decompresed_KeyFrame

def compress_not_KeyFrame(Frame_class: data, KeyFrame: data, KeyFrameDecompressed:data, dzielnik, encode):
    Compress_data = data()
    ## TO DO
    Frame_class.Y = (Frame_class.Y - KeyFrameDecompressed.Y)//dzielnik
    Frame_class.Cb = (Frame_class.Cb - KeyFrameDecompressed.Cb)//dzielnik
    Frame_class.Cr = (Frame_class.Cr - KeyFrameDecompressed.Cr)//dzielnik

    Compress_data.Y = encode(Frame_class.Y)
    Compress_data.Cb = encode(Frame_class.Cb)
    Compress_data.Cr = encode(Frame_class.Cr)
    return Compress_data

def decompress_not_KeyFrame(Compress_data: data,  KeyFrame: data, KeyFrameDecompressed:data, dzielnik, subsampling, decode):
    Y=decode(Compress_data.Y)
    Cb=decode(Compress_data.Cb)
    Cr=decode(Compress_data.Cr)
    ## TO DO
    
    Y = Y*dzielnik + KeyFrameDecompressed.Y
    Cb = Cb*dzielnik + KeyFrameDecompressed.Cb
    Cr = Cr*dzielnik + KeyFrameDecompressed.Cr
    return frame_layers_to_image(Y,Cr,Cb, subsampling)

def plotDiffrence(ReferenceFrame, DecompressedFrame, roi):

    ReferenceFrameRGB = cv2.cvtColor(ReferenceFrame.astype(np.uint8), cv2.COLOR_YCrCb2RGB)
    DecompressedFrameRGB = cv2.cvtColor(DecompressedFrame.astype(np.uint8), cv2.COLOR_YCrCb2RGB)
    diffRGB = ReferenceFrameRGB[roi[0]:roi[1], roi[2]:roi[3]].astype(float) - DecompressedFrameRGB[roi[0]:roi[1], roi[2]:roi[3]].astype(float)
    diffYCrCb=ReferenceFrame[roi[0]:roi[1], roi[2]:roi[3]].astype(float) - DecompressedFrame[roi[0]:roi[1], roi[2]:roi[3]].astype(float)

    fig, axs = plt.subplots(4, 3, sharey=True)
    fig.tight_layout()
    fig.set_size_inches(8,9)
    
    axs[0, 0].set_title("ref RGB")
    axs[0, 0].imshow(ReferenceFrameRGB[roi[0]:roi[1], roi[2]:roi[3]])

    axs[0, 1].set_title("diff  RGB")
    axs[0, 1].imshow(diffRGB, vmin=np.min(diffRGB), vmax=np.max(diffRGB))

    axs[0, 2].set_title("dec RGB")
    axs[0, 2].imshow(DecompressedFrameRGB[roi[0]:roi[1], roi[2]:roi[3]])

    
    axs[1, 0].set_title("ref Y")
    axs[1, 0].imshow(ReferenceFrame[roi[0]:roi[1], roi[2]:roi[3], 0], cmap=plt.cm.gray)

    axs[1, 1].set_title("diff Y")
    axs[1, 1].imshow(diffYCrCb[:,:,0], vmin=np.min(diffYCrCb[:,:,0]), vmax=np.max(diffYCrCb[:,:,0]))

    axs[1, 2].set_title("dec Y")
    axs[1, 2].imshow(DecompressedFrame[roi[0]:roi[1], roi[2]:roi[3], 0], cmap=plt.cm.gray)


    axs[2, 0].set_title("ref Cr")
    axs[2, 0].imshow(ReferenceFrame[roi[0]:roi[1], roi[2]:roi[3], 1], cmap=plt.cm.gray)

    axs[2, 1].set_title("diff Cr")
    axs[2, 1].imshow(diffYCrCb[:,:,1], vmin=np.min(diffYCrCb[:,:,1]), vmax=np.max(diffYCrCb[:,:,1]))

    axs[2, 2].set_title("dec Cr")
    axs[2, 2].imshow(DecompressedFrame[roi[0]:roi[1], roi[2]:roi[3], 1], cmap=plt.cm.gray)


    axs[3, 0].set_title("ref Cb")
    axs[3, 0].imshow(ReferenceFrame[roi[0]:roi[1], roi[2]:roi[3], 2], cmap=plt.cm.gray)

    axs[3, 1].set_title("diff Cb")
    axs[3, 1].imshow(diffYCrCb[:,:,2], vmin=np.min(diffYCrCb[:,:,2]), vmax=np.max(diffYCrCb[:,:,2]))

    axs[3, 2].set_title("dec Cb")
    axs[3, 2].imshow(DecompressedFrame[roi[0]:roi[1], roi[2]:roi[3], 2], cmap=plt.cm.gray)

    return fig, axs

##############################################################################
####     Głowna pętla programu      ##########################################
##############################################################################
def main_funtion(encoding, decoding, plik, ile, auto_pause_frames, wyswietlaj_klatki, plot_frames, ROI, div, frame_counter, sub):
    document.add_heading(f"File:{plik}, subsampling={sub}, divider={div}, KeyFrame={frame_counter}", 2)
    cap = cv2.VideoCapture(os.path.join(INPUT_DIR,plik))

    if ile<0:
        ile=int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

    if wyswietlaj_klatki:
        cv2.namedWindow('Normal Frame')
        cv2.namedWindow('Decompressed Frame')

    compression_information=np.zeros((3,ile))

    for i in range(ile):
        ret, frame = cap.read()
        if wyswietlaj_klatki:
            cv2.imshow('Normal Frame',frame)

        frame=cv2.cvtColor(frame,cv2.COLOR_BGR2YCrCb)
        Frame_class = frame_image_to_class(frame, sub)

        if (i % frame_counter) == 0: # pobieranie klatek kluczowych
            KeyFrame = compress_KeyFrame(Frame_class, encode = encoding)
            cY=KeyFrame.Y
            cCb=KeyFrame.Cb
            cCr=KeyFrame.Cr
            Decompresed_Frame, Decompresed_KeyFrame = decompress_KeyFrame(KeyFrame, subsampling=sub, decode = decoding)

        else: # kompresja
            Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, Decompresed_KeyFrame, dzielnik=div, encode = encoding)
            cY=Compress_data.Y
            cCb=Compress_data.Cb
            cCr=Compress_data.Cr
            Decompresed_Frame = decompress_not_KeyFrame(Compress_data,  KeyFrame, Decompresed_KeyFrame, dzielnik=div, subsampling=sub, decode=decoding)
        
        compression_information[0, i] = (frame[:,:,0].size - cY.size)/frame[:,:,0].size
        compression_information[1, i] = (frame[:,:,0].size - cCb.size)/frame[:,:,0].size
        compression_information[2, i] = (frame[:,:,0].size - cCr.size)/frame[:,:,0].size

        if wyswietlaj_klatki:
            cv2.imshow('Decompressed Frame',cv2.cvtColor(Decompresed_Frame,cv2.COLOR_YCrCb2BGR))
        
        if np.any(plot_frames==i): # rysuj wykresy
            for r in ROI:
                fig, axs = plotDiffrence(frame, Decompresed_Frame, r)
                # document.add_paragraph(f"ROI = {r}, frame number = {i}")
                fig.suptitle(f"ROI = {r}, frame number = {i}")

                memfile = BytesIO()
                fig.tight_layout()
                fig.savefig(memfile)
                document.add_picture(memfile)
                memfile.close()
            
        if np.any(auto_pause_frames==i):
            cv2.waitKey(-1) #wait until any key is pressed
        
        k = cv2.waitKey(1) & 0xff
        
        if k==ord('q'):
            break
        elif k == ord('p'):
            cv2.waitKey(-1) #wait until any key is pressed

    if encoding != encode_none:

        fig, axs = plt.subplots(3,1)
        axs[0].plot(np.arange(0,ile), compression_information[0,:]*100)
        axs[0].set_title("Y size")

        axs[1].plot(np.arange(0,ile), compression_information[1,:]*100)
        axs[1].set_title("Cb size")

        axs[2].plot(np.arange(0,ile), compression_information[2,:]*100)
        axs[2].set_title("Cr size")

        fig.tight_layout()
        fig.suptitle(f"File:{plik}, subsampling={sub}, divider={div}, KeyFrame={frame_counter}")

        memfile = BytesIO()
        fig.tight_layout()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

        # plt.show()



document = Document()
for section in document.sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
document.add_heading('Lab8 - Milosz Zubala (zm49455)', 0)

if __name__ == "__main__":

    encoding = encode_ByteRun
    decoding = decode_ByteRun                         # katalog z plikami wideo
    plik="clip_1.mp4"                       # nazwa pliku
    ile=50                                  # ile klatek odtworzyć? <0 - całość
    auto_pause_frames=np.array([])          # automatycznie za pauzuj dla klatki
    wyswietlaj_klatki=True                 # czy program ma wyświetlać klatki
    plot_frames=np.array([30, 45])          # automatycznie wyrysuj wykresy
    ROI = [[200, 300, 830, 930]]            # wyświetlane fragmenty (można podać kilka )

    SUBSAMPLINGS = ["4:1:0","4:2:2", "4:2:0", "4:1:1", "4:4:0", "4:4:4"]
    DIVISIONS = [4, 8, 16]
    KEY_FRAMES = [4, 8, 16]

    document.add_heading(f"Badanie jakości dla różnych parametrów bez użycia RLE lub ByteRun", 1)
    for div in DIVISIONS:
        for frame_counter in KEY_FRAMES:
            for sub in SUBSAMPLINGS:      
                main_funtion(encoding = encode_none,
                            decoding = decode_none,
                            plik = plik,
                            ile = ile,
                            auto_pause_frames = auto_pause_frames,
                            wyswietlaj_klatki = wyswietlaj_klatki,
                            plot_frames = plot_frames,
                            ROI = ROI,
                            div = div,
                            frame_counter = frame_counter,
                            sub = sub)

    document.add_heading(f"Badanie skuteczności kompresji z użyciem RLE lub ByteRun (0.4 pkt)", 1)

    document.add_heading(f"Obserwacje", 1)
    document.add_paragraph(f"placeholder")

    document.save(os.path.join(OUTPUT_DIR,'lab8_milosz_zubala.docx'))
    os.system(os.path.join(OUTPUT_DIR,'lab8_milosz_zubala.docx'))