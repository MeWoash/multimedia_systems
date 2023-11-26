import numpy as np
import os
import sys
import matplotlib.pyplot as plt
import pandas as pd

#doc generation
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

INPUT_DIR = os.path.normpath(os.path.join(__file__,'..','input'))

TEST1 = np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]).astype(int)
TEST2 = np.array([1,2,3,1,2,3,1,2,3]).astype(int)
TEST3 = np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]).astype(int)
TEST4 = np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]).astype(int)
TEST5 = np.zeros((1,520)).astype(int)
TEST6 = np.arange(0,521,1).astype(int)
TEST7 = np.eye(7).astype(int)
TEST8 = np.dstack([np.eye(7),np.eye(7),np.eye(7)]).astype(int)
TEST9 = np.ones((1,1,1,1,1,1,10)).astype(int)

TEST_SUITE = [
    TEST1,
    TEST2,
    TEST3,
    TEST4,
    TEST5,
    TEST6,
    TEST7,
    TEST8,
    TEST9,
]


def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

def index_same_numbers(start, data) -> list:
    maxsize = data.size
    element_to_find = data[start]
    repetitions = 1
    current_index = start + 1

    while current_index < maxsize:
        if data[current_index] == element_to_find:
            repetitions+=1
        else:
            break
        current_index += 1

    return [int(repetitions), int(current_index)]

def index_different_numbers(start:int, data:int) -> list:
    maxsize = data.size
    last_element = data[start]
    streak = 1
    current_index = start + 1

    while current_index < maxsize:
        if data[current_index] != last_element:
            streak+=1
            last_element = data[current_index]
        else:
            break
        current_index += 1

    if current_index < maxsize:
        return [int(streak-1), int(current_index-1)]
    else :
        return [int(streak), int(current_index)]
    
def encode_RLE(data: np.ndarray):
    data_copy = data.copy()
    max_size = data.size
    x=np.array([len(data.shape)])
    x=np.concatenate([x, data.shape])

    data_copy = data_copy.flatten()
    bufor = np.zeros(int(np.prod(data.shape)*2,)).astype(int)

    start = 0
    bufor_current_index = 0

    while start < max_size:
        current_element = data_copy[start]

        repetitions, start = index_same_numbers(start, data_copy)
        bufor[[bufor_current_index, bufor_current_index+1]] = [repetitions, current_element]
        bufor_current_index += 2
        

    bufor = np.concatenate([x, bufor[:bufor_current_index]])
    return bufor

def decode_RLE(data: np.ndarray):
    max_size = data.size
    shape = data[1:int(data[0]+1)].astype(int)
    bufor_size = np.sum(data[int(data[0])+1::2])
    bufor = np.zeros(int(bufor_size)).astype(int)

    start = int(data[0])+1
    bufor_current_index = 0

    while start<max_size:
        repetitions = data[start]
        element = data[start + 1]
        seq = np.repeat(element, repetitions)

        bufor[bufor_current_index:bufor_current_index+int(repetitions)] = seq
        bufor_current_index += int(repetitions)
        start+=2

    bufor=bufor.reshape(shape)

    return bufor

def encode_ByteRun(data: np.ndarray):
    data_copy = data.copy()
    max_size = data.size
    x=np.array([len(data.shape)])
    x=np.concatenate([x, data.shape])

    data_copy = data_copy.flatten()
    bufor = np.zeros(int(np.prod(data.shape)*2,)).astype(int)

    start = 0
    bufor_current_index = 0

    while start < max_size:
        current_element = data_copy[start]

        repetitions_tmp, start_tmp = index_same_numbers(start, data_copy)
        if repetitions_tmp > 1:
            bufor[[bufor_current_index, bufor_current_index+1]] = [-repetitions_tmp+1, current_element]
            repetitions, start = repetitions_tmp, start_tmp
            bufor_current_index += 2
        else:
            repetitions_tmp, start_tmp = index_different_numbers(start, data_copy)
            bufor[bufor_current_index : bufor_current_index+repetitions_tmp+1] = np.concatenate([np.array([repetitions_tmp-1]), data_copy[start:start_tmp]])

            repetitions, start = repetitions_tmp, start_tmp
            bufor_current_index += repetitions + 1

    bufor = np.concatenate([x, bufor[:bufor_current_index]])
    return bufor

def decode_ByteRun(data: np.ndarray):
    max_size = data.size
    shape = data[1:int(data[0]+1)].astype(int)
    bufor_size = np.prod(shape)
    bufor = np.zeros(int(bufor_size)).astype(int)

    start = int(data[0])+1
    bufor_current_index = 0

    while start<max_size:
        repetitions = data[start]

        if repetitions <= -1:
            repetitions = -repetitions + 1
            element = data[start + 1]
            seq = np.repeat(element, repetitions)

            bufor[bufor_current_index:bufor_current_index+int(repetitions)] = seq
            bufor_current_index += int(repetitions)
            start+=2
        else:
            repetitions = repetitions + 1
            bufor[bufor_current_index:bufor_current_index+repetitions] = data[start+1:start+repetitions+1]
            bufor_current_index += int(repetitions)
            start+=repetitions+1
        

    bufor=bufor.reshape(shape)

    return bufor

def imgToUInt8(img):
    
    if np.issubdtype(img.dtype, np.floating):
        img = (img*255).astype(int)
    return img   

if __name__ == "__main__":
    document = Document()
    for section in document.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    document.add_heading('Lab5 - Milosz Zubala (zm49455)', 0)

    document.add_heading('Data payload', 1)
    document.add_paragraph(f"Dane są kompresowane do 1-wymiarowego obiektu np.ndarray. Na początku tablicy jest załączona informacja o ilości wymiarów, a za nią ich rozmiary.")

    document.add_heading('Tests', 1)

    document.add_heading('RLE', 2)
    for index, test_vector in enumerate(TEST_SUITE):

        document.add_heading(f'Test {index+1}', 3)

        encoded = encode_RLE(test_vector)
        decoded = decode_RLE(encoded)
        equal = np.all(test_vector == decoded)
        assert(equal)

        before_compr = get_size(test_vector)
        after_compr = get_size(encoded)
        cr = np.abs(before_compr/after_compr)
        pr = np.abs(after_compr/before_compr)*100

        document.add_heading(f"Test Vector", 4)
        document.add_paragraph(str(test_vector))

        document.add_heading(f"Encoded Vector", 4)
        document.add_paragraph(str(encoded))

        document.add_heading(f"Decoded Vector", 4)
        document.add_paragraph(f"{decoded}")

        document.add_heading("Metrics", 4)
        document.add_paragraph(f"Are equal = {equal}")
        document.add_paragraph(f"Before compression = {before_compr}")
        document.add_paragraph(f"After compression = {after_compr}")
        document.add_paragraph(f"CR = {cr}")
        document.add_paragraph(f"PR = {pr}%")


    document.add_heading('ByteRun', 2)
    for index, test_vector in enumerate(TEST_SUITE):

        document.add_heading(f'Test {index+1}', 3)

        encoded = encode_ByteRun(test_vector)
        decoded = decode_ByteRun(encoded)
        equal = np.all(test_vector == decoded)
        assert(equal)

        before_compr = get_size(test_vector)
        after_compr = get_size(encoded)
        cr = np.abs(before_compr/after_compr)
        pr = np.abs(after_compr/before_compr)*100

        document.add_heading(f"Test Vector", 4)
        document.add_paragraph(str(test_vector))

        document.add_heading(f"Encoded Vector", 4)
        document.add_paragraph(str(encoded))

        document.add_heading(f"Decoded Vector", 4)
        document.add_paragraph(f"{decoded}")

        document.add_heading("Metrics", 4)
        document.add_paragraph(f"Are equal = {equal}")
        document.add_paragraph(f"Before compression = {before_compr}")
        document.add_paragraph(f"After compression = {after_compr}")
        document.add_paragraph(f"CR = {cr}")
        document.add_paragraph(f"PR = {pr}%")

    PHOTOS = ["dokument.png", "kolorowe.png", "techniczny.png"]

    document.add_heading(f'Photos', 1)

    document.add_heading(f'RLE', 2)
    for index, photo_name in enumerate(PHOTOS):

        document.add_heading(f'Photo {index+1}', 3)
        current_image = imgToUInt8(plt.imread(os.path.join(INPUT_DIR, photo_name)))
        
        encoded = encode_RLE(current_image)
        decoded = decode_RLE(encoded)
        equal = np.all(current_image == decoded)
        assert(equal)

        before_compr = get_size(current_image)
        after_compr = get_size(encoded)
        cr = np.abs(before_compr/after_compr)
        pr = np.abs(after_compr/before_compr)*100
        uniq = len(np.unique(current_image.flatten()))
        shape = current_image.shape

        fig, ax = plt.subplots(1,2)

        ax[0].imshow(current_image)
        ax[1].imshow(decoded)

        fig.suptitle(photo_name)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

        # document.add_heading(f"Data", 4)
        # document.add_paragraph(str(current_image))

        document.add_heading(f"Encoded Vector", 4)
        document.add_paragraph(str(encoded))

        # document.add_heading(f"Decoded Vector", 4)
        # document.add_paragraph(f"{decoded}")

        document.add_heading("Metrics", 4)
        document.add_paragraph(f"Are equal = {equal}")
        document.add_paragraph(f"Before compression = {before_compr}")
        document.add_paragraph(f"After compression = {after_compr}")
        document.add_paragraph(f"CR = {cr}")
        document.add_paragraph(f"PR = {pr}%")
        document.add_paragraph(f"Shape = {shape}")
        document.add_paragraph(f"Unique values = {uniq}")

        fig, ax = plt.subplots(1,1)

        d = current_image.flatten()
        ax.hist(d, bins = uniq)

        fig.suptitle("Values")
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

    document.add_heading(f'ByteRun', 2)
    for index, photo_name in enumerate(PHOTOS):

        document.add_heading(f'Photo {index+1}', 3)
        current_image = imgToUInt8(plt.imread(os.path.join(INPUT_DIR, photo_name)))
        
        encoded = encode_ByteRun(current_image)
        decoded = decode_ByteRun(encoded)
        equal = np.all(current_image == decoded)
        assert(equal)

        before_compr = get_size(current_image)
        after_compr = get_size(encoded)
        cr = np.abs(before_compr/after_compr)
        pr = np.abs(after_compr/before_compr)*100
        uniq = len(np.unique(current_image.flatten()))
        shape = current_image.shape

        fig, ax = plt.subplots(1,2)

        ax[0].imshow(current_image)
        ax[1].imshow(decoded)

        fig.suptitle(photo_name)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

        # document.add_heading(f"Data", 4)
        # document.add_paragraph(str(current_image))

        document.add_heading(f"Encoded Vector", 4)
        document.add_paragraph(str(encoded))

        # document.add_heading(f"Decoded Vector", 4)
        # document.add_paragraph(f"{decoded}")

        document.add_heading("Metrics", 4)
        document.add_paragraph(f"Are equal = {equal}")
        document.add_paragraph(f"Before compression = {before_compr}")
        document.add_paragraph(f"After compression = {after_compr}")
        document.add_paragraph(f"CR = {cr}")
        document.add_paragraph(f"PR = {pr}%")
        document.add_paragraph(f"Shape = {shape}")
        document.add_paragraph(f"Unique values = {uniq}")

        fig, ax = plt.subplots(1,1)

        d = current_image.flatten()
        ax.hist(d, bins = uniq)

        fig.suptitle("Values")
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

document.add_heading(f'Obserwacje', 1)
document.add_paragraph(f"Kompresja RLE dla ciągu danych, które się różnią na sąsiadujących pozycjach może zabierać nawet 2 razy więcej pamięci. ByteRun naprawia ten problem poprzez możliwość przepisania n różnych bitów. Kompresje najlepiej radzą sobie dla jednolitego tła o tym samym kolorze - wzzór dokumentu.Próba kodowania ciągle zmieniających się bajtów ostatecznie zabiera więcej miejsca.")

document.save(os.path.normpath(os.path.join(__file__,'..','output','lab5_milosz_zubala.docx')))



