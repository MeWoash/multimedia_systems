import numpy as np
import os
import sys
import matplotlib.pyplot as plt
import pandas as pd

from scipy.interpolate import interp1d

#doc generation
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

#sound
import sounddevice as sd
import soundfile as sf
import scipy.fftpack

INPUT_DIR = os.path.normpath(os.path.join(__file__,'..','input'))
OUTPUT_DIR = os.path.normpath(os.path.join(__file__,'..','output'))

def plotAudio(Signal: np.ndarray, Fs:int, TimeMargin=[0, 0.02], fsize = None, axs = None) -> None:
    
    if fsize is None:
        fsize = 2**10

    if axs is None:
        fig, axs = plt.subplots(2,1)
        fig.tight_layout()
    else:
        fig = axs[0].get_figure()

    xTime = np.arange(0, Signal.shape[0])/Fs
    axs[0].plot(xTime, Signal)
    axs[0].set_xlim(TimeMargin)
    axs[0].set_xlabel('s')

    yf = scipy.fftpack.fft(Signal, fsize)

    xFreq = np.arange(0, Fs/2, Fs/fsize)
    ydB = 20*np.log10( np.abs(yf[:fsize//2]))

    axs[1].plot(xFreq, ydB)
    axs[1].set_xlabel("Hz")
    axs[1].set_ylabel('dB')

    ymax = np.argmax(ydB)
    xmax = xFreq[ymax]
    axs[1].axvline(xmax, linestyle="dotted", color='r')

    return fig, axs

def quant(data: np.ndarray, bit: int):
    data_copy = np.copy(data)
    d=float(2**bit-1)
    typ=data_copy.dtype
    #if float
    if np.issubdtype(data_copy.dtype, np.floating):
        m = -1.0
        n = 1.0
    #if int
    else:
        m = np.iinfo(data_copy.dtype).min
        n = np.iinfo(data_copy.dtype).max

    DataF = data.astype(float)
    DataF = ((DataF-m)/(n-m))
    DataF = np.round(DataF*d)
    DataF = ((DataF/d)*(n-m))+m

    # kwantyzacja na DataF
    return DataF.astype(typ)

def no_quant(data: np.ndarray, bit: int):
    return data


def A_Law_encode(x: np.ndarray, A:float = 87.6, **kwargs):
    x_out = np.copy(x)
    mask1 = np.abs(x_out)<(1/A)
    mask2 = np.logical_not(mask1)

    x_out[mask1] = (A*np.abs(x_out[mask1])) / (1+np.log(A))
    x_out[mask2] = (1+np.log(A*np.abs(x_out[mask2]))) / (1 + np.log(A))
    x_out = x_out * np.sign(x)

    return x_out


def A_Law_decode(y: np.ndarray, A:float = 87.6, **kwargs):
    y_out = np.copy(y)
    mask1 = np.abs(y_out) < 1/(1+np.log(A))
    mask2 = np.logical_not(mask1)

    y_out[mask1] = (np.abs(y_out[mask1])*(1+np.log(A)))/(A)
    y_out[mask2] = np.exp(np.abs(y_out[mask2])*(1+np.log(A))-1) / (A)

    y_out = y_out * np.sign(y)

    return y_out


def Mu_Law_encode(x: np.ndarray, mu:float=255, **kwargs):
    x_out = np.copy(x)

    x_out = np.log(1+mu*np.abs(x_out))/np.log(1+mu)

    x_out = x_out * np.sign(x)
    return x_out

def Mu_Law_decode(y: np.ndarray, mu:float=255, **kwargs):
    y_out = np.copy(y)

    y_out = (1/mu)*(np.power(1+mu, np.abs(y_out))-1)

    y_out = y_out * np.sign(y)

    return y_out

def DPCM_encode(x, bit, **kwargs):
    y=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=quant(x[i]-e,bit)
        e+=y[i]
    return y

def DPCM_decode(Y, **kwargs):
    X=np.zeros(Y.shape)
    e=0
    for i in range(0, Y.shape[0]):
        if i-1 < 0:
            e = 0
        else:
            e = X[i-1]
        X[i]=Y[i]+e
    return X

def DPCM_encode_pred(x, bit, predictor, n, **kwargs): 
    y=np.zeros(x.shape)
    xp=np.zeros(x.shape)
    e=0
    for i in range(1,x.shape[0]):
        y[i]=quant(x[i]-e,bit)
        xp[i]=y[i]+e
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        e=predictor(xp[idx])
    return y

def DPCM_decode_pred(Y, predictor, n, **kwargs): 
    X=np.zeros(Y.shape)
    xp=np.zeros(Y.shape)
    e=0
    for i in range(1,Y.shape[0]):
        X[i]=Y[i]+e
        xp[i]=X[i]
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        e=predictor(xp[idx])

    return X

def no_pred(X):
    return X[-1]

def mean_pred(X:np.ndarray):
    mean = np.mean(X)
    # print(f"V = {X}, mean = {mean}")
    return mean

def sound_load(file_path: str):
    data, fs = sf.read(file_path)
    return data, fs

def sound_play(data: np.ndarray, fs:int):
    sd.play(data, fs)

if __name__ == "__main__":

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    document.add_heading('Lab6 - Milosz Zubala (zm49455)', 0)

    #TESTING
    x1=np.linspace(-1,1,1000)

    x2=np.linspace(-1,1,3000)
    y2=0.9*np.sin(np.pi*x2*4)

    TEST_VECTOR = [ #x, y, encode, decode, nbits
                    [x1, x1, A_Law_encode, A_Law_decode, 8],
                    [x1, x1, Mu_Law_encode, Mu_Law_decode, 8],
                    [x2, y2, A_Law_encode, A_Law_decode, 8],
                    [x2, y2, Mu_Law_encode, Mu_Law_decode, 8],
                ]
    
    document.add_heading(f"Testing", 1)

    for x_test, y_test, encode_f, decode_f, bits in TEST_VECTOR:
        document.add_heading(f"Methods: {encode_f.__name__}, {decode_f.__name__}, bits: {bits}",2)


        encoded = encode_f(y_test)
        encoded_quant = quant(encoded, bit = bits)
        decoded = decode_f(encoded_quant)

        fig, axs = plt.subplots(1,2,figsize = (8,3))
        axs[0].set_title(f"Compression {encode_f.__name__}")
        axs[0].plot(x_test, y_test, label="Original")
        axs[0].plot(x_test, encoded_quant, linestyle="-", label="Compr quant")
        axs[0].plot(x_test, encoded, linestyle="-", label="Compr")
        axs[0].legend()

        axs[1].set_title(f"Decompression {decode_f.__name__}")
        axs[1].plot(x_test, y_test, label="Original")
        axs[1].plot(x_test, decoded, linestyle="-", label="Decompr")
        axs[1].legend()

        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

    TEST_VECTOR2 = [#x, y, encode, decode, pred, n, nbits, xlim, ylim
        [x2, y2, DPCM_encode, DPCM_decode, None, None, 6, [-0.5, -0.25], [0,1]],
        [x2, y2, DPCM_encode_pred, DPCM_decode_pred, mean_pred, 5, 6, [-0.5, -0.25], [0,1]]
        ]
    for x_test, y_test, encode_f, decode_f, pred, n, bits, xlim, ylim in TEST_VECTOR2:
        
        if pred is not None:
            document.add_heading(f"Methods: {encode_f.__name__}, {decode_f.__name__}, bits: {bits}, predictor: {pred.__name__}, n: {n}",2)
        else:
            document.add_heading(f"Methods: {encode_f.__name__}, {decode_f.__name__}, bits: {bits}")
        encoded = encode_f(y_test, bit=bits, predictor=pred, n=n)
        decoded = decode_f(encoded, predictor=pred, n=n)

        fig, axs = plt.subplots(4,1,figsize = (8,7))
        axs[0].set_title(f"Compression {encode_f.__name__}")
        axs[0].plot(x_test, y_test, label="Original")
        axs[0].plot(x_test, encoded, linestyle="-", label="Compr quant")
        axs[0].legend()

        axs[1].set_title(f"Compression {encode_f.__name__}")
        axs[1].plot(x_test, y_test, label="Original")
        axs[1].plot(x_test, encoded, linestyle="-", label="Compr quant")
        axs[1].legend()
        axs[1].set_xlim(xlim)
        axs[1].set_ylim(ylim)

        axs[2].set_title(f"Decompression {decode_f.__name__}")
        axs[2].plot(x_test, y_test, label="Original")
        axs[2].plot(x_test, decoded, linestyle="-", label="Decompr")
        axs[2].legend()

        axs[3].set_title(f"Decompression {decode_f.__name__}")
        axs[3].plot(x_test, y_test, label="Original")
        axs[3].plot(x_test, decoded, linestyle="-", label="Decompr")
        axs[3].legend()
        axs[3].set_xlim(xlim)
        axs[3].set_ylim(ylim)

        fig.tight_layout()
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()

    FILES = ["sing_low1.wav", "sing_medium1.wav", "sing_high1.wav"]
    BITS = [7,6,5,4,3,2]
    METHODS = [
        [A_Law_encode, A_Law_decode, quant, {}],
        [Mu_Law_encode, Mu_Law_decode, quant, {}],
        [DPCM_encode, DPCM_decode, no_quant, {}],
        [DPCM_encode_pred, DPCM_decode_pred, no_quant, {"predictor":mean_pred, "n":5}],
        ] 

    document.add_heading(f"Odsłuch",1)
    for file in FILES:
        file_path = os.path.join(INPUT_DIR, file)
        data, fs = sound_load(file_path=file_path)

        sf.write(os.path.join(OUTPUT_DIR, f"{file}_original.wav"),
                 data,
                 fs)

        document.add_heading(f"Plik {file}",2)
        for encode_f, decode_f, quant_f, params in METHODS:
            
            document.add_heading(f"Metody: {encode_f.__name__}, {decode_f.__name__}",3)
            for bits in BITS:
                # encoded = encode_f(data, bit = bits, **params)
                # encoded_quant = quant(encoded, bit = bits)
                # decoded = decode_f(encoded_quant, **params)

                document.add_paragraph(f"{bits}bits: ")
                # fname = f"{file}_{encode_f.__name__}_{bits}bits.wav"
                # sf.write(os.path.join(OUTPUT_DIR, fname),
                #  decoded,
                #  fs)
                # print(f"Created {fname}")

    # document.save(os.path.normpath(os.path.join(__file__,'..','output','lab6_milosz_zubala.docx')))
    plt.show()



