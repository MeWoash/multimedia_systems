import numpy as np
import os
import sys
import matplotlib.pyplot as plt
import pandas as pd

import cv2

#doc generation
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

import scipy.fftpack

INPUT_DIR = os.path.normpath(os.path.join(__file__,'..','input'))
OUTPUT_DIR = os.path.normpath(os.path.join(__file__,'..','output'))
QUANTS = {
    "QY":np.array([
            [16, 11, 10, 16, 24,  40,  51,  61],
            [12, 12, 14, 19, 26,  58,  60,  55],
            [14, 13, 16, 24, 40,  57,  69,  56],
            [14, 17, 22, 29, 51,  87,  80,  62],
            [18, 22, 37, 56, 68,  109, 103, 77],
            [24, 36, 55, 64, 81,  104, 113, 92],
            [49, 64, 78, 87, 103, 121, 120, 101],
            [72, 92, 95, 98, 112, 100, 103, 99],
            ]),

    "QC":np.array([
            [17, 18, 24, 47, 99, 99, 99, 99],
            [18, 21, 26, 66, 99, 99, 99, 99],
            [24, 26, 56, 99, 99, 99, 99, 99],
            [47, 66, 99, 99, 99, 99, 99, 99],
            [99, 99, 99, 99, 99, 99, 99, 99],
            [99, 99, 99, 99, 99, 99, 99, 99],
            [99, 99, 99, 99, 99, 99, 99, 99],
            [99, 99, 99, 99, 99, 99, 99, 99],
            ]),

    "QN":np.ones((8,8))
}

def index_same_numbers(start, data) -> list:
    maxsize = data.size
    element_to_find = data[start]
    repetitions = 1
    current_index = start + 1

    while current_index < maxsize:
        if data[current_index] == element_to_find:
            repetitions+=1
        else:
            break
        current_index += 1

    return [int(repetitions), int(current_index)]

def index_different_numbers(start:int, data:int) -> list:
    maxsize = data.size
    last_element = data[start]
    streak = 1
    current_index = start + 1

    while current_index < maxsize:
        if data[current_index] != last_element:
            streak+=1
            last_element = data[current_index]
        else:
            break
        current_index += 1

    if current_index < maxsize:
        return [int(streak-1), int(current_index-1)]
    else :
        return [int(streak), int(current_index)]

def encode_ByteRun(data: np.ndarray):
    data_copy = data.copy()
    max_size = data.size
    x=np.array([len(data.shape)])
    x=np.concatenate([x, data.shape])

    data_copy = data_copy.flatten()
    bufor = np.zeros(int(np.prod(data.shape)*2,)).astype(int)

    start = 0
    bufor_current_index = 0

    while start < max_size:
        current_element = data_copy[start]

        repetitions_tmp, start_tmp = index_same_numbers(start, data_copy)
        if repetitions_tmp > 1:
            bufor[[bufor_current_index, bufor_current_index+1]] = [-repetitions_tmp+1, current_element]
            repetitions, start = repetitions_tmp, start_tmp
            bufor_current_index += 2
        else:
            repetitions_tmp, start_tmp = index_different_numbers(start, data_copy)
            bufor[bufor_current_index : bufor_current_index+repetitions_tmp+1] = np.concatenate([np.array([repetitions_tmp-1]), data_copy[start:start_tmp]])

            repetitions, start = repetitions_tmp, start_tmp
            bufor_current_index += repetitions + 1

    bufor = np.concatenate([x, bufor[:bufor_current_index]])
    return bufor

def decode_ByteRun(data: np.ndarray):
    max_size = data.size
    shape = data[1:int(data[0]+1)].astype(int)
    bufor_size = np.prod(shape)
    bufor = np.zeros(int(bufor_size)).astype(int)

    start = int(data[0])+1
    bufor_current_index = 0

    while start<max_size:
        repetitions = data[start]

        if repetitions <= -1:
            repetitions = -repetitions + 1
            element = data[start + 1]
            seq = np.repeat(element, repetitions)

            bufor[bufor_current_index:bufor_current_index+int(repetitions)] = seq
            bufor_current_index += int(repetitions)
            start+=2
        else:
            repetitions = repetitions + 1
            bufor[bufor_current_index:bufor_current_index+repetitions] = data[start+1:start+repetitions+1]
            bufor_current_index += int(repetitions)
            start+=repetitions+1
        

    bufor=bufor.reshape(shape)

    return bufor

def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

def dct2(a):
    return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
    return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

class JPEG_container:
    def __init__(self, Y, Cb, Cr, OGShape, Ratio="4:4:4", QY=np.ones((8,8)), QC=np.ones((8,8))):
        self.shape = OGShape
        self.Y=Y
        self.Cb=Cb
        self.Cr=Cr
        self.ChromaRatio=Ratio
        self.QY=QY
        self.QC=QC

def imgToUInt8(img):
    if np.issubdtype(img.dtype, np.floating):
        img = (img*255).astype('uint8')   
    return img

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B

def CompressBlock(block,Q):
    #DCT
    block = block - 128
    block = dct2(block)
    
    #QUANT
    block=np.round(block/Q).astype(int)

    #ZIGZAG
    vector = zigzag(block)

    return vector

def DecompressBlock(vector,Q):
    #ZIGZAG
    block = zigzag(vector)

    #QUANT
    block=block*Q

    #IDCT
    block = idct2(block)
    block = block + 128
    return block

## podział na bloki
# L - warstwa kompresowana
# S - wektor wyjściowy
def CompressLayer(L,Q):
    S=np.array([])
    for w in range(0,L.shape[0],8):
        for k in range(0,L.shape[1],8):
            block=L[w:(w+8),k:(k+8)]
            S=np.append(S, CompressBlock(block,Q))
    return S

## wyodrębnianie bloków z wektora 
# L - warstwa o oczekiwanym rozmiarze
# S - długi wektor zawierający skompresowane dane
def DecompressLayer(S,Q):
    L=np.zeros( (128, 128) )
    for idx,i in enumerate(range(0,S.shape[0],64)):
        vector=S[i:(i+64)]
        m=L.shape[0]/8
        k=int((idx%m)*8)
        w=int((idx//m)*8)
        L[w:(w+8),k:(k+8)]=DecompressBlock(vector,Q)
    return L

def chroma_subsampling(A, Ratio):
    B = np.copy(A)
    if Ratio == "4:2:2":
        B = B[:,::2]
    elif Ratio == "4:2:0":
        pass
    else:
        pass
    return B


def chroma_resampling(A, Ratio):
    B = np.copy(A)
    if Ratio == "4:2:2":
        B = B[B!=0].reshape(128,64)
        B = np.repeat(B, 2, axis=1).reshape(128,128)
    elif Ratio == "4:2:0":
        pass
    else:
        pass
    return B

def CompressJPEG(RGB, Ratio="4:4:4", QY=np.ones((8,8)), QC=np.ones((8,8))):
    # RGB -> YCrCb
    YCrCb=cv2.cvtColor(RGB,cv2.COLOR_RGB2YCrCb).astype(int)

    # JPEG = verX(...); zapisać dane z wejścia do kalsy
    JPEG = JPEG_container(Y=YCrCb[:,:,0],
                Cr=YCrCb[:,:,1],
                Cb=YCrCb[:,:,2],
                OGShape=YCrCb.shape,
                Ratio=Ratio,
                QY=QY,
                QC=QC)
    
    # Tu chroma subsampling
    JPEG.Cr = chroma_subsampling(JPEG.Cr, JPEG.ChromaRatio)
    JPEG.Cb = chroma_subsampling(JPEG.Cb, JPEG.ChromaRatio)

    #Kompresja stratna
    JPEG.Y=CompressLayer(JPEG.Y,JPEG.QY)
    JPEG.Cr=CompressLayer(JPEG.Cr,JPEG.QC)
    JPEG.Cb=CompressLayer(JPEG.Cb,JPEG.QC)
    
    # tu dochodzi kompresja bezstratna
    JPEG.Y=encode_ByteRun(JPEG.Y)
    JPEG.Cr=encode_ByteRun(JPEG.Cr)
    JPEG.Cb=encode_ByteRun(JPEG.Cb)

    return JPEG

def DecompressJPEG(JPEG):

    # dekompresja bezstratna
    Y=decode_ByteRun(JPEG.Y)
    Cr=decode_ByteRun(JPEG.Cr)
    Cb=decode_ByteRun(JPEG.Cb)

    #deKompresja stratna
    Y=DecompressLayer(Y, JPEG.QY)
    Cr=DecompressLayer(Cr, JPEG.QC)
    Cb=DecompressLayer(Cb, JPEG.QC)

    # Tu chroma resampling
    Cr = chroma_resampling(Cr, JPEG.ChromaRatio)
    Cb = chroma_resampling(Cb, JPEG.ChromaRatio)

    # tu rekonstrukcja obrazu
    YCrCb=np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)

    # YCrCb -> RGB
    RGB=cv2.cvtColor(YCrCb.astype(np.uint8),cv2.COLOR_YCrCb2RGB)

    return RGB

if __name__ == "__main__":

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    document.add_heading('Lab7 - Milosz Zubala (zm49455)', 0)

    file = "photo1.jpg"
    file_path = os.path.join(INPUT_DIR, file)
    img = plt.imread(file_path)
    img = imgToUInt8(img)
    
    document.add_heading(f"Original photo", 1)
    fig, axs = plt.subplots(1, 1)
    fig.suptitle("Original photo")
    fig.tight_layout()
    axs.imshow(img)

    memfile = BytesIO()
    fig.savefig(memfile)
    document.add_picture(memfile)
    memfile.close()

    document.add_heading(f"Compression", 1)

    bounds = [#x, y, dx, dy
            [[150, 270], [128, 128]],
            [[400,190], [128, 128]],
            [[580,100], [128, 128]]
              ]
    ratios = ["4:4:4",
              "4:2:2"]
    quants= [#Qy, Qc
            ["QN", "QN"],
            ["QY", "QC"],
            ]

    for bound in bounds:
        x = bound[0][1]
        y = bound[0][0]
        dx = bound[1][1]
        dy = bound[1][0]

        before = np.copy(img)[x:x+dx, y:y+dx,:]
        # document.add_heading(f"x0 y0: {bound[0]}, dx dy: {bound[1]}", 2)

        for ratio in ratios:
            # document.add_heading(f"Ratio: {ratio}", 3)

            for quant in quants:
                qy_name, qc_name = quant

                # document.add_heading(f"Qy: {qy_name}, Qc: {qc_name}", 4)
                compressed = CompressJPEG(before, Ratio=ratio, QY = QUANTS[qy_name], QC = QUANTS[qc_name])
                decompressed = DecompressJPEG(compressed)

                fig, axs = plt.subplots(4, 2, sharey=True)
                fig.set_size_inches(8,10)

                fig.suptitle(f"Ratio: {ratio}, Qy: {qy_name}, Qc: {qc_name}")

                # obraz oryginalny 
                axs[0,0].set_title("Before RGB")
                axs[0,0].imshow(before) #RGB 

                # jego warstwy w Y Cr Cb -> dopisać konwersję
                before_YCrCb=cv2.cvtColor(before, cv2.COLOR_RGB2YCrCb).astype(int)
                before_Y = before_YCrCb[:,:,0]
                before_Cr = before_YCrCb[:,:,1]
                before_Cb = before_YCrCb[:,:,2]

                axs[1,0].set_title("Before Y")
                axs[1,0].imshow(before_Y, cmap=plt.cm.gray)

                axs[2,0].set_title("Before Cr")
                axs[2,0].imshow(before_Cr, cmap=plt.cm.gray)

                axs[3,0].set_title("Before Cb")
                axs[3,0].imshow(before_Cb, cmap=plt.cm.gray)

                # obraz po dekompresji
                axs[0,1].set_title("After RGB")
                axs[0,1].imshow(decompressed) #RGB

                # jego warstwy w Y Cr Cb -> dopisać konwersję
                after_YCrCb=cv2.cvtColor(decompressed, cv2.COLOR_RGB2YCrCb).astype(int)
                after_Y = after_YCrCb[:,:,0]
                after_Cr = after_YCrCb[:,:,1]
                after_Cb = after_YCrCb[:,:,2]

                axs[1,1].set_title("After Y")
                axs[1,1].imshow(after_Y, cmap=plt.cm.gray)

                axs[2,1].set_title("After Cr")
                axs[2,1].imshow(after_Cr, cmap=plt.cm.gray)

                axs[3,1].set_title("After Cb")
                axs[3,1].imshow(after_Cb, cmap=plt.cm.gray)

                memfile = BytesIO()
                fig.tight_layout()
                fig.savefig(memfile)
                document.add_picture(memfile)
                memfile.close()

                size_before_Y = get_size(before_Y)
                size_before_Cr = get_size(before_Cr)
                size_before_Cb = get_size(before_Cb)
                size_sum_before_YCrCb = size_before_Y + size_before_Cr + size_before_Cb

                size_compressed_Y = get_size(compressed.Y)
                size_compressed_Cr = get_size(compressed.Cr)
                size_compressed_Cb = get_size(compressed.Cb)
                size_sum_compressed_YCrCb = size_compressed_Y + size_compressed_Cr + size_compressed_Cb

                cr = np.abs(size_sum_before_YCrCb/size_sum_compressed_YCrCb)
                pr = np.abs(size_sum_compressed_YCrCb/size_sum_before_YCrCb)*100

                document.add_paragraph(f"""
                Y before: {size_before_Y}, after: {size_compressed_Y}
                Cr before: {size_before_Cr}, after: {size_compressed_Cr}
                Cb before: {size_before_Cb}, after: {size_compressed_Cb}
                Sum before: {size_sum_before_YCrCb}, after: {size_sum_compressed_YCrCb}
                Cr={cr}, Pr={pr}%
                """)
                
                
    document.add_heading(f'Obserwacje', 1)
    document.add_paragraph(f"Dla kodowania ByteRun subsampling z ratio 4:2:2 a 4:4:4 i braku kwantyzacji zmniejsza pamięć około o 30%. Przy włączonej kwantyzacji oszczędność to zaledwie około 4%. Dla Ratio 4:4:4 włączenie kwantyzacji zaoszczędza około 70% pamięci, a dla 4:2:2 około 50%. Wniosek jest taki, że najwięcej pamięci zaoszczędzimy z subsamplingiem 4:2:2 i włączoną kwantyzacją. Dla włączonej kwantyzacji nie było widać żadnych różnic w obrazkach. Artefakty były widoczne po dodaniu subsamplingu 4:2:2")

    document.save(os.path.join(OUTPUT_DIR,'lab7_milosz_zubala.docx'))
    os.system(os.path.join(OUTPUT_DIR,'lab7_milosz_zubala.docx'))
    # plt.show()