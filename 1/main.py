import numpy as np
import matplotlib.pyplot as plt
import sounddevice as sd
import soundfile as sf
import scipy.fftpack
from io import BytesIO
from docx import Document
from docx.shared import Inches

import sys
import os
 
CURRENFT_FILE_DIR = os.path.join(__file__,'..')
sys.path.append(f'{CURRENFT_FILE_DIR}/../common')
# from functions import *


def zad1():
    # 1
    data, fs = sf.read(f'{CURRENFT_FILE_DIR}/input/sound1.wav', dtype='float32')

    #a)
    sound_L = data[:,0]
    sound_R = data[:,1]
    sound_mix = data.mean(axis=1)
    print(sound_mix)

    sf.write(f'{CURRENFT_FILE_DIR}/output/sound_L.wav', sound_L, fs)
    sf.write(f'{CURRENFT_FILE_DIR}/output/sound_R.wav', sound_R, fs)
    sf.write(f'{CURRENFT_FILE_DIR}/output/sound_mix.wav', sound_mix, fs)

    #b)
    plt.subplot(2,1,1)
    x = np.arange(0, 1/fs*data.shape[0], 1/fs)
    plt.plot(x, data[:,0])

    #c)
    data, fs = sf.read(f'{CURRENFT_FILE_DIR}/input/sin_440Hz.wav', dtype=np.int32)

    plt.figure()
    plt.subplot(2,1,1)
    plt.plot(np.arange(0,data.shape[0])/fs, data)

    plt.subplot(2,1,2)
    yf = scipy.fftpack.fft(data)
    plt.plot(np.arange(0,fs, 1.0*fs/(yf.size)),np.abs(yf))

    #d)
    fsize=2**8

    plt.figure()
    plt.subplot(2,1,1)
    plt.plot(np.arange(0,data.shape[0])/fs,data)
    plt.subplot(2,1,2)
    yf = scipy.fftpack.fft(data, fsize)
    plt.plot(np.arange(0, fs/2, fs/fsize), 20*np.log10( np.abs(yf[:fsize//2])))

def plotAudio(Signal: np.ndarray, Fs:int, TimeMargin=[0, 0.02], fsize = None, axs = None) -> None:
    
    if fsize is None:
        fsize = Signal.shape[0]

    if axs is None:
        fig, axs = plt.subplots(2,1)
        fig.tight_layout()

    xTime = np.arange(0, Signal.shape[0])/Fs
    axs[0].plot(xTime, Signal)
    axs[0].set_xlim(TimeMargin)
    axs[0].set_xlabel('s')

    yf = scipy.fftpack.fft(Signal, fsize)

    xFreq = np.arange(0, Fs/2, Fs/fsize)
    ydB = 20*np.log10( np.abs(yf[:fsize//2]))

    axs[1].plot(xFreq, ydB)
    axs[1].set_xlabel("Hz")
    axs[1].set_ylabel('dB')

    ymax = np.argmax(ydB)
    xmax = xFreq[ymax]
    axs[1].axvline(xmax, linestyle="dotted", color='r')

    return xmax

    

def zad2():

    document = Document()
    document.add_heading('Lab1',0) # tworzenie nagłówków druga wartość to poziom nagłówka 


    files=['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
    Margins=[[0,0.02],[0.133,0.155]]
    for file in files:
        document.add_heading(f'Plik - {file}',2)
        for i,Margin in enumerate(Margins):
            document.add_heading(f'Time margin {Margin}',3) # nagłówek sekcji, mozę być poziom wyżej
            fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
        
            ############################################################
            # Tu wykonujesz jakieś funkcje i rysujesz wykresy

            data, fs = sf.read(f'{CURRENFT_FILE_DIR}/input/{file}', dtype='float32')
            fmax = plotAudio(data, fs, axs = axs)
            
            ############################################################
            
            fig.suptitle(f'Time margin {Margin}') # Tytuł wykresu
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora 
            
        
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
            
            memfile.close()
            ############################################################
            # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
            document.add_paragraph(f'częstotliwość widma z najwyższą wartością = {fmax}')
            ############################################################

    document.save(f'{CURRENFT_FILE_DIR}/output/report_zad2.docx') # zapis do pliku


def zad3():
    document = Document()
    document.add_heading('Lab1',0) # tworzenie nagłówków druga wartość to poziom nagłówka 


    files=['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
    fsizes=[2**8,2**12,2**16]
    for file in files:
        document.add_heading('Plik - {}'.format(file),2)
        for i,fsize in enumerate(fsizes):
            document.add_heading(f'fsize = {fsize}',3) # nagłówek sekcji, mozę być poziom wyżej
            fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
        
            ############################################################
            # Tu wykonujesz jakieś funkcje i rysujesz wykresy

            data, fs = sf.read(f'{CURRENFT_FILE_DIR}/input/{file}', dtype='float32')
            fmax = plotAudio(data, fs, axs = axs, fsize = fsize)
            
            ############################################################
            
            fig.suptitle(f'fsize = {fsize}') # Tytuł wykresu
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora 
            
        
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
            
            memfile.close()
            ############################################################
            # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
            document.add_paragraph(f'częstotliwość widma z najwyższą wartością = {fmax}')
            ############################################################

    document.save(f'{CURRENFT_FILE_DIR}/output/report_zad3.docx') # zapis do pliku


if __name__ == "__main__":
    zad1()
    # plt.show()
    zad2()
    zad3()
   


