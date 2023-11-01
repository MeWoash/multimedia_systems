import numpy as np
import os
import sys
import matplotlib.pyplot as plt
import pandas as pd

#doc generation
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

INPUT_DIR = os.path.normpath(os.path.join(__file__,'..','input'))

PALETT8 = np.array(
[
    [0.0, 0.0, 0.0,],
    [0.0, 0.0, 1.0,],
    [0.0, 1.0, 0.0,],
    [0.0, 1.0, 1.0,],
    [1.0, 0.0, 0.0,],
    [1.0, 0.0, 1.0,],
    [1.0, 1.0, 0.0,],
    [1.0, 1.0, 1.0,],
])

PALETT16 =  np.array(
[
    [0.0, 0.0, 0.0,], 
    [0.0, 1.0, 1.0,],
    [0.0, 0.0, 1.0,],
    [1.0, 0.0, 1.0,],
    [0.0, 0.5, 0.0,], 
    [0.5, 0.5, 0.5,],
    [0.0, 1.0, 0.0,],
    [0.5, 0.0, 0.0,],
    [0.0, 0.0, 0.5,],
    [0.5, 0.5, 0.0,],
    [0.5, 0.0, 0.5,],
    [1.0, 0.0, 0.0,],
    [0.75, 0.75, 0.75,],
    [0.0, 0.5, 0.5,],
    [1.0, 1.0, 1.0,], 
    [1.0, 1.0, 0.0,]
])

M1 = np.array(
[
    [0, 2],
    [3, 1]
]
)

M2 = np.array(
[
    [0, 8, 2, 10],
    [12, 4, 14, 6],
    [3, 11, 1, 9],
    [15, 7, 13, 5]
]
)

def calc_Mpre(M: np.ndarray, n: int)->np.ndarray:
    Mpre = (M+1) / (2*n)**2 - 0.5
    return Mpre

def gen_M(_2n: int):
    
    if _2n == 0:
        return np.array([[0]])
    elif _2n == 1:
        return M1
    elif _2n == 2:
        return M2
    else:
        n = _2n//2
        coeff = (_2n)**2
        M_prev = gen_M(n)

        A = coeff * M_prev
        B = coeff * M_prev + 2
        C = coeff * M_prev + 3
        D = coeff * M_prev + 1

        out_M = np.block([[A, B],
                        [C, D]])
        return out_M

def get_text_width(document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000

def imgToFloat(img):
    if np.issubdtype(img.dtype, np.unsignedinteger):
        img = img/255.0
    return img    

def calcY1(img):
    Y1 = 0.299 * img[:,:,0] + 0.587 * img[:,:,1] + 0.114 * img[:,:,2]
    return Y1

def colorFit(pixel: np.ndarray, Pallet:np.ndarray) -> np.ndarray:
    closestColor = Pallet[np.argmin(np.linalg.norm(pixel - Pallet, axis=1))].squeeze()
    # print(f"Pixel {pixel} -> {closestColor} in palett: {Pallet}")
    return np.copy(closestColor)


def kwant_colorFit(img:np.ndarray, Pallet:np.ndarray) -> np.ndarray:
    out_img = img.copy()

    img_shape = out_img.shape
    for w in range(img_shape[0]):
        for k in range(img_shape[1]):
            out_img[w,k] = colorFit(img[w,k], Pallet)

    return out_img



def show_photos(photo_data: list[np.ndarray],
                photo_title: list[str],
                greyscale: bool,
                figsize = (8,3)) -> [plt.Figure, plt.Axes]:
    assert(len(photo_data)==len(photo_title))

    n_axes = len(photo_data)
    fig, axs = plt.subplots(1,len(photo_data), figsize=figsize, layout = 'compressed')
    if n_axes==1:
         axs = [axs]
    for i in range(n_axes):
        if greyscale:
            axs[i].imshow(photo_data[i], cmap=plt.cm.gray)
        else:
            axs[i].imshow(photo_data[i])
        axs[i].set_title(photo_title[i])
    return fig, axs

def n_bit_greyscale_pallet(n: int=2):
    bits = 2**n
    return np.linspace(0,1,bits).reshape(bits,1)

def dithering_random(img:np.ndarray, Pallet:np.ndarray) -> np.ndarray:
    r = np.random.rand(*img.shape)
    out_img = np.where(img>=r, 1, 0)

    return out_img

def dithering_ordered(img:np.ndarray, Pallet:np.ndarray, r=1, mapping:np.ndarray = M2, mapping_n = 2) -> np.ndarray:
    out_img = img.copy()
    mpre = calc_Mpre(mapping, mapping_n)
    out_img_shape = out_img.shape

    for i_row in range(out_img_shape[0]):
         for i_col in range(out_img_shape[1]):
            new_pixel = out_img[i_row, i_col] + (r * mpre[i_row % (mapping_n*2), i_col % (mapping_n*2)])
            out_img[i_row, i_col] = colorFit(new_pixel, Pallet)
            # print(f"new_pixel: {new_pixel}, quant: {out_img[i_row, i_col]}")

    return out_img

def dithering_floyd_steinberg(img:np.ndarray, Pallet:np.ndarray) -> np.ndarray:
    out_img = img.copy()
    out_img_shape = out_img.shape

    for i_row in range(out_img_shape[0]):
        for i_col in range(out_img_shape[1]):
        
            old_pixel = np.copy(out_img[i_row, i_col])
            out_img[i_row, i_col] = colorFit(old_pixel, Pallet)
            quant_error = old_pixel - out_img[i_row, i_col]

            if i_col + 1 + 1 < out_img_shape[1]:
                out_img[i_row, i_col + 1] = out_img[i_row, i_col + 1] + quant_error * 7 / 16

            if i_row + 1 < out_img_shape[0] and i_col - 1 < out_img_shape[1]:
                out_img[i_row + 1, i_col - 1] = out_img[i_row + 1, i_col - 1] + quant_error * 3 / 16

            if i_row + 1 < out_img_shape[0]:
                out_img[i_row + 1 ,i_col] = out_img[i_row + 1 ,i_col] + quant_error * 5 / 16

            if i_row + 1 < out_img_shape[0] and i_col + 1 < out_img_shape[1]:
                out_img[i_row + 1, i_col + 1] = out_img[i_row + 1, i_col + 1] + quant_error * 1 / 16

    return out_img


if __name__ == "__main__":

    columns = ["filename", "greyscale", "pallets"]
    data = [
    ['GS_0001.tif', True, {
                            "pallet_1_bit":n_bit_greyscale_pallet(1),
                            "pallet_2_bit":n_bit_greyscale_pallet(2),
                            "pallet_4_bit":n_bit_greyscale_pallet(4)
                            }],

    ['GS_0002.png', True, {
                            "pallet_1_bit":n_bit_greyscale_pallet(1),
                            "pallet_2_bit":n_bit_greyscale_pallet(2),
                            "pallet_4_bit":n_bit_greyscale_pallet(4)
                           }],

    ['GS_0003.png', True, {
                            "pallet_1_bit":n_bit_greyscale_pallet(1),
                            "pallet_2_bit":n_bit_greyscale_pallet(2),
                            "pallet_4_bit":n_bit_greyscale_pallet(4)
                           }],

    ['SMALL_0009.jpg', False, {
                                "pallet_8_bit":PALETT8,
                                "pallet_16_bit":PALETT16
                                }],

    ['SMALL_0007.jpg', False, {
                                "pallet_8_bit":PALETT8,
                                "pallet_16_bit":PALETT16
                                }],

    ['SMALL_0006.jpg', False, {
                                "pallet_8_bit":PALETT8,
                                "pallet_16_bit":PALETT16
                                }],

    ['SMALL_0005.jpg', False, {
                                "pallet_8_bit":PALETT8,
                                "pallet_16_bit":PALETT16
                                }],
    ]
    df = pd.DataFrame(data=data, columns=columns)



    document = Document()
    for section in document.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    document.add_heading('Lab3 - Milosz Zubala (zm49455)', 0)

    document.add_heading(f"Kwantyzacja i dithering",1)
    for index, row in df.iterrows():
        current_image = plt.imread(os.path.join(INPUT_DIR, row['filename']))
        current_image = imgToFloat(current_image)

        document.add_heading(f"Zdjecie: {row['filename']}",2)
        if row['greyscale']==True:
            if len(current_image.shape)>2:
                current_image = calcY1(current_image)
            
            for p_name, pallet in row['pallets'].items():
                
                document.add_heading(f"Paleta {p_name}",3)

                data = [current_image]
                labels = ['original']

                data.append(kwant_colorFit(current_image, pallet))
                labels.append('quantization')

                if len(np.unique(pallet)) == 2:
                    data.append(dithering_random(current_image, pallet))
                    labels.append('dith random')
                
                data.append(dithering_ordered(current_image, pallet))
                labels.append('dith ordered')

                data.append(dithering_floyd_steinberg(current_image, pallet))
                labels.append('dith floyd stteinberg')

                fig,ax = show_photos(data, labels, True)

                fig.suptitle(p_name)
                memfile = BytesIO()
                fig.savefig(memfile)
                
                document.add_picture(memfile)
                memfile.close()
            
        else:
            
            for p_name, pallet in row['pallets'].items():
                
                document.add_heading(f"Paleta {p_name}",3)

                data = [current_image]
                labels = ['original']

                data.append(kwant_colorFit(current_image, pallet))
                labels.append('quantization')

                data.append(dithering_ordered(current_image, pallet))
                labels.append('dith ordered')

                data.append(dithering_floyd_steinberg(current_image, pallet))
                labels.append('dith floyd stteinberg')

                fig,ax = show_photos(data, labels, True)

                fig.suptitle(p_name)
                memfile = BytesIO()
                fig.savefig(memfile)
                
                document.add_picture(memfile)
                memfile.close()

    document.add_heading(f"Obserwacje",1)

    document.save(os.path.normpath(os.path.join(__file__,'..','output','lab3_milosz_zubala.docx')))
    plt.show()