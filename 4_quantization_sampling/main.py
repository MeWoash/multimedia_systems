import numpy as np
import os
import sys
import matplotlib.pyplot as plt
import pandas as pd

from scipy.interpolate import interp1d

#doc generation
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

#sound
import sounddevice as sd
import soundfile as sf
import scipy.fftpack

INPUT_DIR = os.path.normpath(os.path.join(__file__,'..','input'))

def plotAudio(Signal: np.ndarray, Fs:int, TimeMargin=[0, 0.02], fsize = None, axs = None) -> None:
    
    if fsize is None:
        fsize = 2**10

    if axs is None:
        fig, axs = plt.subplots(2,1)
        fig.tight_layout()
    else:
        fig = axs[0].get_figure()

    xTime = np.arange(0, Signal.shape[0])/Fs
    axs[0].plot(xTime, Signal)
    axs[0].set_xlim(TimeMargin)
    axs[0].set_xlabel('s')

    yf = scipy.fftpack.fft(Signal, fsize)

    xFreq = np.arange(0, Fs/2, Fs/fsize)
    ydB = 20*np.log10( np.abs(yf[:fsize//2]))

    axs[1].plot(xFreq, ydB)
    axs[1].set_xlabel("Hz")
    axs[1].set_ylabel('dB')

    ymax = np.argmax(ydB)
    xmax = xFreq[ymax]
    axs[1].axvline(xmax, linestyle="dotted", color='r')

    return fig, axs

def quant(data, bit):
    d=float(2**bit-1)
    typ=data.dtype
    #if float
    if np.issubdtype(data.dtype, np.floating):
        m = -1.0
        n = 1.0
    #if int
    else:
        m = np.iinfo(data.dtype).min
        n = np.iinfo(data.dtype).max

    DataF = data.astype(float)
    DataF = ((DataF-m)/(n-m))
    DataF = np.round(DataF*d)
    DataF = ((DataF/d)*(n-m))+m

    # kwantyzacja na DataF
    return DataF.astype(typ)

def decim(data: np.ndarray, fs: int, n) -> [np.ndarray, int]:
    return np.copy(data[::n]), fs//n

def interp(data: np.ndarray, fs_old:int, fs_new:int):
    N=data.size
    N1= int(N//fs_old * fs_new)

    x=np.linspace(0,N,N)
    x1=np.linspace(0,N,N1)

    t=np.linspace(0,N/fs,N)
    t1=np.linspace(0,N/fs,N1)

    metode_lin=interp1d(x,data)
    metode_nonlin=interp1d(x,data, kind='cubic')

    y_lin=metode_lin(x1).astype(data.dtype)
    y_nonlin=metode_nonlin(x1).astype(data.dtype)

    return y_lin, y_nonlin

def sound_load(file_path: str):
    data, fs = sf.read(file_path)
    return data, fs

def sound_play(data: np.ndarray, fs:int):
    print("Playing sound...")
    sd.play(data, fs)
    inp = input("Finished playing!, type your input:\n")
    return inp

COLUMNS = ["filename"]
DATA = [
    ['sin_60Hz.wav'],
    ['sin_440Hz.wav'],
    ['sin_8000Hz.wav'],
    ['sin_combined.wav'],
]

COLUMNS_LISTEN = ["filename"]
DATA_LISTEN = [
    ['sing_high1.wav'],
    ['sing_medium1.wav'],
    ['sing_low1.wav'],
]

if __name__ == "__main__":
    
    df = pd.DataFrame(data=DATA, columns=COLUMNS)
    df_listen = pd.DataFrame(data=DATA_LISTEN, columns=COLUMNS_LISTEN)

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    document.add_heading('Lab4 - Milosz Zubala (zm49455)', 0)

    document.add_heading(f"Kwantyzacja",1)

    document.add_heading(f"Quant test", 2)
    lin = np.linspace(-1,1,20)
    lin_quant = quant(lin, 2)
    
    fig, axs = plt.subplots(1,2)
    fig.tight_layout()
    axs[0].plot(lin, lin)
    axs[1].plot(lin, lin_quant)
    fig.suptitle(f"Quantization test")
    memfile = BytesIO()
    fig.savefig(memfile)
    document.add_picture(memfile)
    memfile.close()
    document.add_page_break()

    for index, row in df.iterrows():
        document.add_heading(f"Plik {row['filename']}",2)

        file_path = os.path.join(INPUT_DIR, row['filename'])
        data, fs = sound_load(file_path=file_path)
        
        fig, axs = plotAudio(data, fs)
        fig.suptitle(f"Original: {row['filename']}")
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile)
        memfile.close()
        document.add_page_break()

        for bit in [4, 8, 16, 24]:
            document.add_heading(f"bits: {bit}",3)
            data_quant = quant(data, bit)

            fig, axs = plotAudio(data_quant, fs)
            fig.suptitle(f"{row['filename']}, bits:{bit}")
            memfile = BytesIO()
            fig.savefig(memfile)
            document.add_picture(memfile)
            memfile.close()
            document.add_page_break()

        document.add_heading(f"Decymacja",1)
        for n in [2,4,6,10,24]:
            document.add_heading(f"decymacja: {n}",3)

            data_dec, fs_new = decim(data, fs, n)

            fig, axs = plotAudio(data_dec, fs_new)
            fig.suptitle(f"{row['filename']}, n:{n}")
            memfile = BytesIO()
            fig.savefig(memfile)
            document.add_picture(memfile)
            memfile.close()
            document.add_page_break()

        document.add_heading(f"Interpolacja",1)
        for fs_new in [2000,4000,8000,11999,16000,16953,24000,41000]:
            document.add_heading(f"czestotliwosc: {fs_new}",3)

            data_interp_lin, data_interp_nonlin = interp(data, fs, fs_new)

            fig, axs = plotAudio(data_interp_lin, fs_new)
            fig.suptitle(f"{row['filename']}, linear, fs:{fs_new}")
            memfile = BytesIO()
            fig.savefig(memfile)
            document.add_picture(memfile)
            memfile.close()

            fig, axs = plotAudio(data_interp_nonlin, fs_new)
            fig.suptitle(f"{row['filename']}, nonlinear, fs:{fs_new}")
            memfile = BytesIO()
            fig.savefig(memfile)
            document.add_picture(memfile)
            memfile.close()
            document.add_page_break()

    document.add_heading(f"Odsluch",1)
    for index, row in df_listen.iterrows():

        file_path = os.path.join(INPUT_DIR, row['filename'])
        data, fs = sound_load(file_path=file_path)

        document.add_heading(f"Kwantyzacja",2)
        for bit in [4, 8]:
            document.add_heading(f"bits: {bit}",3)
            data_quant = quant(data, bit)

            print(f"Quantization: Playing {row['filename']}, bit: {bit}")
            inp = sound_play(data_quant, fs)
            document.add_paragraph(inp)

        document.add_heading(f"Decymacja",2)
        for n in [4,6,10,24]:
            document.add_heading(f"n: {n}",3)
            data_dec, fs_new = decim(data, fs, n)

            print(f"Decimation: Playing {row['filename']}, n: {n}")
            inp = sound_play(data_dec, fs_new)
            document.add_paragraph(inp)

        document.add_heading(f"Interpolacja",2)
        for fs_new in [4000, 8000, 11999, 16000, 16953]:
            document.add_heading(f"czestotliwosc: {fs_new}",3)
            data_interp_lin, data_interp_nonlin = interp(data, fs, fs_new)  

            document.add_heading(f"liniowa:",4)
            print(f"Interpolation lin: Playing {row['filename']}, fs: {fs_new}")
            inp = sound_play(data_interp_lin, fs_new)
            document.add_paragraph(inp)

            document.add_heading(f"nieliniowa:",4)
            print(f"Interpolation nonlin: Playing {row['filename']}, fs: {fs_new}")
            inp = sound_play(data_interp_lin, fs_new)
            document.add_paragraph(inp)

    # plt.show()
    document.add_heading(f"Obserwacje",1)
    document.save(os.path.normpath(os.path.join(__file__,'..','output','lab4_milosz_zubala.docx')))
