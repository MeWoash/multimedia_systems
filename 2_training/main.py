from io import BytesIO
from docx import Document
from docx.shared import Inches

import numpy as np
import matplotlib.pyplot as plt
import cv2

import sys
import os

import pandas as pd
 
CURRENFT_FILE_DIR = os.path.join(__file__,'..')
sys.path.append(f'{CURRENFT_FILE_DIR}/../common')
# from functions import *

imgA1 = plt.imread(f'{CURRENFT_FILE_DIR}/input/A1.png')
imgA2 = plt.imread(f'{CURRENFT_FILE_DIR}/input/A2.jpg')
imgA3 = plt.imread(f'{CURRENFT_FILE_DIR}/input/A3.png')
imgA4 = plt.imread(f'{CURRENFT_FILE_DIR}/input/A4.jpg')
imgB1 = plt.imread(f'{CURRENFT_FILE_DIR}/input/B01.png')
imgB2 = plt.imread(f'{CURRENFT_FILE_DIR}/input/B02.jpg')



def test():
    img1 = plt.imread(f'{CURRENFT_FILE_DIR}/input/A1.png')

    print(img1.dtype)
    print(img1.shape)
    print(np.min(img1),np.max(img1))

def imgToUInt8(img):
    
    if np.issubdtype(img.dtype, np.floating):
        img_new = (img*255).astype('uint8')
        
    return img_new   

def calcY1(img):
    Y1 = 0.299 * img[:,:,0] + 0.587 * img[:,:,1] + 0.114 * img[:,:,2]
    return Y1


def calcY2(img):
    Y2 = 0.2126 * img[:,:,0] + 0.7152 * img[:,:,1] + 0.0722 * img[:,:,2]
    return Y2

def imgToFloat(img):

    if np.issubdtype(img.dtype, np.unsignedinteger):
        img = img/255.0

    return img    



def zad1():
    
    R=imgA1[:,:,0]
    plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=1)

    Y2=calcY2(imgA1)
    plt.imshow(Y2, cmap=plt.cm.gray, vmin=0, vmax=1)

    # cv2.imshow('a', cv2.cvtColor(imgA1, cv2.COLOR_RGB2BGR))

def zad2():
    zad2_sub(imgB1)

def zad2_sub(img):
    img_copy = img.copy()

    print(f'Typ obrazu: {img_copy.dtype}')
    print(f'Rozmiar obrazu: {img_copy.shape}')
    print(f'min: {np.min(img_copy)}, max: {np.max(img_copy)}')

    fig, axs = plt.subplots(3,3)
    fig.tight_layout()

    axs[0,0].imshow(img_copy)
    axs[0,0].set_title('O')

    Y1=calcY1(img_copy)
    axs[0,1].imshow(Y1, cmap=plt.cm.gray)
    axs[0,1].set_title('Y1')

    Y2=calcY2(img_copy)
    axs[0,2].imshow(Y2, cmap=plt.cm.gray)
    axs[0,2].set_title('Y2')


    axs[1,0].imshow(img_copy[:,:,0], cmap=plt.cm.gray)
    axs[1,0].set_title('R')

    axs[1,1].imshow(img_copy[:,:,1], cmap=plt.cm.gray)
    axs[1,1].set_title('G')

    axs[1,2].imshow(img_copy[:,:,2], cmap=plt.cm.gray)
    axs[1,2].set_title('B')

    r_only = img_copy.copy()
    r_only[:,:,1:]=0
    axs[2,0].imshow(r_only, cmap=plt.cm.gray)
    axs[2,0].set_title('R')

    g_only = img_copy.copy()
    g_only[:,:,0]=0
    g_only[:,:,2]=0
    axs[2,1].imshow(g_only, cmap=plt.cm.gray)
    axs[2,1].set_title('G')

    b_only = img_copy.copy()
    b_only[:,:,0:2]=0
    axs[2,2].imshow(b_only, cmap=plt.cm.gray)
    axs[2,2].set_title('B')

    return fig,axs



def zad3():
    df = pd.DataFrame()

    df = pd.DataFrame(data={'Filename':[f'{CURRENFT_FILE_DIR}/input/B02.jpg'],
                            'Grayscale':[False],
                            'Fragments':[[[300,20,500,220],[200,200,400,400],[400,400,600,600]]]
                        })

    document = Document()
    document.add_heading('Lab2 - Milosz Zubala (zm49455)',0) # tworzenie nagłówków druga wartość to poziom nagłówka 

    for index, row in df.iterrows():
        img = plt.imread(row['Filename'])

        document.add_heading(f"Plik - {row['Filename']}",2)

        if row['Grayscale']:
            pass
            # GS image - teraz teraz nas nie intersuje
        else:
            # Obraz kolowowy
            if row['Fragments'] is not None:
                # mamy nie pustą listę fragmentów
                for f in row['Fragments']:
                    fragment = img[f[0]:f[2],f[1]:f[3]].copy()
                    fig, axs = zad2_sub(fragment)

                    fig.suptitle(f'Wycinek = {f}')
                    fig.tight_layout(pad=1.5) # poprawa czytelności 
                    memfile = BytesIO() # tworzenie bufora
                    fig.savefig(memfile) # z zapis do bufora 
            
                    document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
                    memfile.close()


    document.save(f'{CURRENFT_FILE_DIR}/output/report_zad3.docx') # zapis do pliku

if __name__ == "__main__":

    # test()
    # zad1()
    # zad2()
    zad3()
    # plt.show()